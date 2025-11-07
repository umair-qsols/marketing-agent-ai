from pathlib import Path
from typing import List, Dict
import streamlit as st

# ---------- NEW DOCX HANDLING ----------
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO
import markdown2
from bs4 import BeautifulSoup

# ---------- LangChain ----------
from langchain_community.document_loaders import Docx2txtLoader, DirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough, RunnableLambda

# ------------------------------------------------------------------
# 1. Configuration
# ------------------------------------------------------------------
EMBEDDING_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
CHROMA_PATH = Path("./chroma_db")
TEMPLATES_DIR = Path("./templates")

# ------------------------------------------------------------------
# 2. Build Vectorstore (now loads .docx via Docx2txtLoader)
# ------------------------------------------------------------------
def build_vectorstore() -> Chroma:
    embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL)

    if CHROMA_PATH.exists():
        print(f"Loading existing vectorstore from {CHROMA_PATH}")
        return Chroma(persist_directory=str(CHROMA_PATH), embedding_function=embeddings)

    print(f"Building new vectorstore from {TEMPLATES_DIR}")

    loader = DirectoryLoader(
        str(TEMPLATES_DIR),
        glob="**/*.docx",
        loader_cls=Docx2txtLoader,               # <-- .docx loader
    )
    docs = loader.load()

    if not docs:
        raise ValueError(f"No .docx files found in {TEMPLATES_DIR}")

    print(f"Loaded {len(docs)} documents")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    splits = splitter.split_documents(docs)

    # ---- enrich metadata ----
    for i, doc in enumerate(splits):
        src = Path(doc.metadata["source"])
        doc.metadata["template_name"] = src.stem
        doc.metadata["doc_type"] = "template"  # not "Template" or mixed case

        # chunk order – Docx2txtLoader does not provide page numbers
        doc.metadata["chunk_order"] = i

        content_lower = doc.page_content.lower()
        if any(m in content_lower for m in ['<company name>', '<client name>', '<date>', '[', 'list:']):
            doc.metadata["doc_type"] = "template"
        else:
            doc.metadata["doc_type"] = "example"

    print(f"Created {len(splits)} chunks")

    vectorstore = Chroma.from_documents(
        documents=splits,
        embedding=embeddings,
        persist_directory=str(CHROMA_PATH),
    )
    print(f"Vectorstore created at {CHROMA_PATH}")
    return vectorstore

vectorstore = build_vectorstore()

# ------------------------------------------------------------------
# 3. LLM Configuration
# ------------------------------------------------------------------
llm = ChatOpenAI(model="gpt-4o-mini", temperature=0.3, api_key=st.secrets["OPENAI_API_KEY"])

# ------------------------------------------------------------------
# 4. Prompt Templates
# ------------------------------------------------------------------
QUESTION_PROMPT = ChatPromptTemplate.from_template(
    """You are an expert marketing consultant analyzing a template document.

Template Name: {template_name}

Template Content:
{context}

Your task is to generate 6-10 specific, actionable questions that will help gather the information needed to complete this template for a client.

IMPORTANT INSTRUCTIONS:
1. Look for placeholders like <COMPANY NAME>, <DATE>, [bracketed instructions], and bullet points with "●" 
2. Identify sections that require specific information (like SWOT analysis, target audience, goals, etc.)
3. Ask questions that directly map to these sections
4. Make questions specific and detailed, not generic
5. Focus on information that would be unique to each client
6. Order questions logically following the template structure

Example good questions:
- "What are the top 3 strengths of your organization that give you a competitive advantage?"
- "Describe your ideal customer in detail - demographics, challenges, and motivations"
- "What are your specific, measurable marketing goals for the next 6-12 months?"

Output ONLY a numbered list of questions:
1. [First question]
2. [Second question]
...

Questions:"""
)

DOC_GEN_PROMPT = ChatPromptTemplate.from_template(
    """You are a professional marketing consultant creating a customized document for a client.

Template Name: {template_name}

ORIGINAL TEMPLATE STRUCTURE:
{template_content}

CLIENT INFORMATION (Answers to questions):
{answers}

YOUR TASK:
1. Use the ORIGINAL TEMPLATE as your exact structural foundation
2. Preserve ALL section headings, formatting, and organization from the template
3. Replace placeholders with client-specific information:
   - <COMPANY NAME> or <CLIENT NAME> → Use the company name from answers
   - <DATE> → Use current date
   - [Bracketed instructions] → Replace with actual content based on answers
   - ● Bullet points → Fill with relevant information from answers

4. Keep the professional tone and style of the original template
5. Expand brief sections with details from the client answers
6. If information is missing for a section, use professional placeholder text like "[To be determined]" or "[Details to be added]"
7. Maintain the same document flow and structure as the template
8. Output in clean **Markdown** format with proper headings (##, ###)

CRITICAL: Your output should look like the completed version of the template, not a completely new document.

Generated Document:"""
)

# ------------------------------------------------------------------
# 5. Helper Functions
# ------------------------------------------------------------------
def get_template_chunks(template_name: str, k: int = 15, doc_type: str = None) -> List:
    """Retrieve chunks for a specific template with proper Chroma filtering."""
    try:
        # Build filter using $and if both conditions are needed
        if doc_type:
            filter_dict = {
                "$and": [
                    {"template_name": {"$eq": template_name}},
                    {"doc_type": {"$eq": doc_type}}
                ]
            }
        else:
            filter_dict = {"template_name": {"$eq": template_name}}

        # Search with filter
        results = vectorstore.similarity_search(
            query="",
            k=k,
            filter=filter_dict
        )

        # Fallback: if no results, try without doc_type
        if not results and doc_type:
            results = vectorstore.similarity_search(
                query="",
                k=k,
                filter={"template_name": {"$eq": template_name}}
            )

        return results

    except Exception as e:
        print(f"Error retrieving chunks: {e}")
        # Manual fallback: get all and filter in Python
        all_docs = vectorstore.similarity_search("", k=200)
        filtered = [
            doc for doc in all_docs
            if doc.metadata.get("template_name") == template_name
            and (doc_type is None or doc.metadata.get("doc_type") == doc_type)
        ]
        return filtered[:k]

def format_docs(docs: List) -> str:
    return "\n\n---\n\n".join(doc.page_content for doc in docs) if docs else "[No content found]"

def get_full_template(template_name: str) -> str:
    docs = get_template_chunks(template_name, k=300)
    if not docs:
        return f"[Template '{template_name}' not found]"
    docs_sorted = sorted(docs, key=lambda d: d.metadata.get("chunk_order", 0))
    return "\n".join(doc.page_content.strip() for doc in docs_sorted)

# ------------------------------------------------------------------
# 6. Runnable Chains
# ------------------------------------------------------------------
def create_question_chain():
    return (
        {
            "template_name": RunnablePassthrough(),
            "context": RunnableLambda(lambda x: format_docs(
                get_template_chunks(x["template_name"], k=10, doc_type="template")
            ))
        }
        | QUESTION_PROMPT
        | llm
        | StrOutputParser()
    )

def create_doc_chain():
    return (
        {
            "template_name": lambda x: x["template_name"],
            "template_content": lambda x: get_full_template(x["template_name"]),
            "answers": lambda x: x["answers"]
        }
        | DOC_GEN_PROMPT
        | llm
        | StrOutputParser()
    )

question_chain = create_question_chain()
doc_chain = create_doc_chain()

# ------------------------------------------------------------------
# 7. Markdown → DOCX conversion
# ------------------------------------------------------------------
def markdown_to_docx(md_text: str) -> BytesIO:
    """
    Convert markdown (produced by the LLM) into a .docx file.
    - Headings become Word styles (Heading 1, Heading 2, …)
    - Bold/italic preserved
    - Bullet/numbered lists preserved
    - Tables are rendered via pandas → docx (fallback to plain text)
    """
    html = markdown2.markdown(md_text, extras=["tables", "fenced-code-blocks"])
    soup = BeautifulSoup(html, "html.parser")

    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    def apply_paragraph(p, elem):
        run = p.add_run(elem.get_text())
        if elem.name == 'strong':
            run.bold = True
        if elem.name == 'em':
            run.italic = True
        if elem.name == 'code':
            run.font.color.rgb = RGBColor(0x0, 0x0, 0x80)

    def process_element(elem, parent_paragraph=None):
        if elem.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(elem.name[1])
            p = doc.add_paragraph(elem.get_text(), style=f'Heading {level}')
            return p
        elif elem.name == 'p':
            p = doc.add_paragraph()
            for child in elem.children:
                if isinstance(child, str):
                    p.add_run(child)
                else:
                    apply_paragraph(p, child)
            return p
        elif elem.name in ['ul', 'ol']:
            for li in elem.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet' if elem.name == 'ul' else 'List Number')
                for child in li.children:
                    if isinstance(child, str):
                        p.add_run(child)
                    else:
                        apply_paragraph(p, child)
            return None
        elif elem.name == 'table':
            rows = elem.find_all('tr')
            if not rows:
                return None
            col_count = len(rows[0].find_all(['td', 'th']))
            table = doc.add_table(rows=len(rows), cols=col_count, style='Table Grid')
            for r_idx, tr in enumerate(rows):
                cells = tr.find_all(['td', 'th'])
                for c_idx, cell in enumerate(cells):
                    table.cell(r_idx, c_idx).text = cell.get_text()
                    if cell.name == 'th':
                        table.cell(r_idx, c_idx).paragraphs[0].runs[0].bold = True
            return None
        else:
            return None

    for child in soup.children:
        if child.name:
            process_element(child)

    # Save to BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# ------------------------------------------------------------------
# 8. Final Document Generation (returns both .docx bytes + markdown)
# ------------------------------------------------------------------
def generate_final_doc(template_name: str, answers: Dict[str, str]) -> tuple[bytes, str]:
    """
    Returns (docx_bytes, markdown_text)
    """
    answers_str = "\n".join(f"Question {k}: {v}" for k, v in answers.items())

    print(f"\nGenerating document from template: {template_name}")
    print(f"Using {len(answers)} answers\n")

    try:
        md_result = doc_chain.invoke({
            "template_name": template_name,
            "answers": answers_str
        })
    except Exception as e:
        md_result = f"# Error Generating Document\n\nError: {str(e)}"

    docx_bio = markdown_to_docx(md_result)
    return docx_bio.read(), md_result